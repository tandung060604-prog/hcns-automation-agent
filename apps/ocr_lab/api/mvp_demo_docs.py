"""Render filled DOCX/PDF artifacts for the MVP leave-request demo.

Uses the frozen template ``hcns format/01_don_xin_nghi_phep_v1.docx`` as the
source of truth for the DOCX export and reuses the same field values for the
PDF export. Both outputs are generated on demand and are synthetic demo data.
"""

from __future__ import annotations

import re
import unicodedata
from io import BytesIO
from pathlib import Path
from typing import Any

# DejaVu covers Vietnamese well and ships with Ubuntu; fall back to Helvetica.
_DEJAVU = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
_DEJAVU_BOLD = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")


def _field(data: dict[str, Any], *names: str) -> str:
    for name in names:
        value = data.get(name)
        if value not in (None, ""):
            return str(value)
    return ""


def _fmt_date(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    match = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", text)
    if match:
        return f"{match.group(3)}/{match.group(2)}/{match.group(1)}"
    return text


def _request_date_parts(value: Any) -> tuple[str, str, str]:
    if value is None:
        return "", "", ""
    text = str(value)
    match = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", text)
    if match:
        return match.group(3), match.group(2), match.group(1)
    return "", "", text


def _replace_dates(text: str, data: dict[str, Any]) -> str:
    start_date = _fmt_date(_field(data, "startDate"))
    end_date = _fmt_date(_field(data, "endDate"))
    expected_return = _fmt_date(_field(data, "expectedReturnDate"))
    if "trở lại" in text or "trơ lại" in text:
        text = text.replace("[dd/mm/yyyy]", expected_return)
    elif "từ ngày" in text or "tu ngay" in text:
        text = text.replace("[dd/mm/yyyy]", start_date, 1)
        text = text.replace("[dd/mm/yyyy]", end_date, 1)
    else:
        text = text.replace("[dd/mm/yyyy]", start_date)
    dd, mm, yyyy = _request_date_parts(_field(data, "requestDate"))
    text = text.replace("[dd]", dd).replace("[mm]", mm).replace("[yyyy]", yyyy)
    return text


def _strip_accent_confound(value: str) -> str:
    """Normalize curly quotes used in the frozen template placeholders."""
    return value.replace("\u2019", "'").replace("\u2018", "'")


def render_leave_docx(data: dict[str, Any], template: Path) -> bytes:
    from docx import Document

    organization = _field(data, "organization", "TÊN CÔNG TY")
    employee_name = _field(data, "employeeName", "Họ và tên")
    job_title = _field(data, "jobTitle")
    department = _field(data, "department")
    address = _field(data, "address")
    phone = _field(data, "phone")
    leave_days = _field(data, "leaveDays", "Số")
    reason = _field(data, "reason")
    handover_to = _field(data, "handoverTo")
    handover_department = _field(data, "handoverDepartment") or department
    handover_tasks = _field(data, "handoverTasks")
    approver = _field(data, "approverName") or employee_name

    document = Document(template)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            original = _strip_accent_confound(run.text)
            if not original or "[" not in original:
                continue
            text = _replace_dates(original, data)
            text = text.replace("[TÊN CÔNG TY]", organization)
            text = text.replace("[Tên công ty]", organization)
            text = text.replace("[Số]", leave_days)
            text = text.replace("[Nêu rõ lý do]", reason)
            text = text.replace("[Mô tả công việc]", handover_tasks)
            text = text.replace("[Chức vụ]", job_title)
            text = text.replace("[Địa chỉ]", address)
            text = text.replace("[Số điện thoại]", phone)
            text = text.replace(
                "[Họ và tên]",
                approver if "PHÊ DUYỆT" in text or "TRƯỞNG BỘ PHẬN" in text else employee_name,
            )
            text = text.replace("[Bộ phận]", handover_department)
            text = text.replace("[Địa danh]", "")
            if text != original:
                run.text = text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        original = _strip_accent_confound(run.text)
                        if not original or "[" not in original:
                            continue
                        text = _replace_dates(original, data)
                        for token, value in [
                            ("[TÊN CÔNG TY]", organization),
                            ("[Tên công ty]", organization),
                            ("[Họ và tên]", employee_name),
                            ("[Chức vụ]", job_title),
                            ("[Bộ phận]", department),
                            ("[Số]", leave_days),
                            ("[Nêu rõ lý do]", reason),
                            ("[Mô tả công việc]", handover_tasks),
                        ]:
                            text = text.replace(token, value)
                        if text != original:
                            run.text = text
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_leave_pdf(data: dict[str, Any]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    if _DEJAVU.is_file():
        pdfmetrics.registerFont(TTFont("DejaVu", str(_DEJAVU)))
        font = "DejaVu"
    else:
        font = "Helvetica"

    organization = _field(data, "organization") or "CÔNG TY ..."
    employee_name = _field(data, "employeeName")
    job_title = _field(data, "jobTitle")
    department = _field(data, "department")
    address = _field(data, "address")
    phone = _field(data, "phone")
    leave_days = _field(data, "leaveDays")
    start_date = _fmt_date(_field(data, "startDate"))
    end_date = _fmt_date(_field(data, "endDate"))
    reason = _field(data, "reason")
    expected_return = _fmt_date(_field(data, "expectedReturnDate"))
    handover_to = _field(data, "handoverTo")
    handover_tasks = _field(data, "handoverTasks")

    title_style = ParagraphStyle(
        "Title",
        fontName=font,
        fontSize=16,
        leading=20,
        alignment=TA_CENTER,
    )
    head_style = ParagraphStyle(
        "Head",
        fontName=font,
        fontSize=10,
        leading=13,
        alignment=TA_CENTER,
    )
    body_style = ParagraphStyle(
        "Body",
        fontName=font,
        fontSize=11,
        leading=18,
    )
    sign_style = ParagraphStyle(
        "Sign",
        fontName=font,
        fontSize=11,
        leading=15,
    )

    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Đơn xin nghỉ phép",
    )

    def p(text: str, style: ParagraphStyle) -> Paragraph:
        return Paragraph(unicodedata.normalize("NFC", text), style)

    content = [
        p(organization, head_style),
        Spacer(1, 10),
        p("ĐƠN XIN NGHỈ PHÉP", title_style),
        Spacer(1, 6),
        p("(Sử dụng bản DOCX của đơn đã nộp khi cần import chữ ký)", head_style),
        Spacer(1, 14),
        p(f"Kính gửi: - Ban Giám đốc {organization}", body_style),
        Spacer(1, 8),
        p(f"Tôi tên là: {employee_name}", body_style),
        p(f"Chức vụ: {job_title}", body_style),
        p(f"Bộ phận: {department}", body_style),
        p(f"Địa chỉ: {address}", body_style),
        p(f"Điện thoại: {phone}", body_style),
        Spacer(1, 8),
        p(
            f"Nay tôi làm đơn này xin nghỉ trong thời gian {leave_days} ngày, "
            f"kể từ ngày {start_date} đến hết ngày {end_date}.",
            body_style,
        ),
        p(f"Lý do xin nghỉ phép: {reason}.", body_style),
        p(f"Tôi dự kiến trở lại làm việc vào ngày {expected_return}.", body_style),
        p(
            f"Tôi đã bàn giao công việc cho: {handover_to} - Bộ phận: {department}.",
            body_style,
        ),
        p(f"Các công việc được bàn giao: {handover_tasks}.", body_style),
        Spacer(1, 22),
        p("TRƯỞNG BỘ PHẬN            NGƯỜI LÀM ĐƠN            PHÊ DUYỆT", sign_style),
        p("[Ký/ghi họ tên]            [Ký/ghi họ tên]            [Ký/ghi họ tên]", sign_style),
    ]
    doc.build(content)
    return output.getvalue()